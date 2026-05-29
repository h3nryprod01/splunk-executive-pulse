"""Generate the Vietnamese audit report (.docx) with embedded screenshots.

    python audit/build_audit.py  ->  audit/Audit_Report_VI.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
SHOTS = HERE / "screenshots"
OUT = HERE / "Audit_Report_VI.docx"


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def main() -> None:
    doc = Document()

    title = doc.add_heading("Báo cáo Audit — Splunk Agentic Ops Hackathon 2026", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Dự án: Splunk Executive Pulse + 2 sản phẩm bổ sung độc lập")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph("Ngày kiểm tra: 2026-05-22  ·  Phạm vi: kiểm thử, giao diện, tuân thủ tài liệu, đề xuất nâng cấp")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Tổng quan
    h(doc, "1. Tổng quan")
    doc.add_paragraph(
        "Repository hiện có 3 sản phẩm độc lập dự thi, nhắm các nhóm giải khác nhau "
        "để không cạnh tranh phiếu lẫn nhau. Tất cả đều chạy được ở chế độ keyless "
        "(không cần API key hay Splunk thật) để demo luôn render."
    )
    rows = [
        ("Sản phẩm", "Track / Giải nhắm", "Vai trò"),
        ("Executive Pulse", "Grand · Best MCP · Best Hosted Models", "Briefing audio 3 phút cho C-suite"),
        ("SPL Copilot", "Best Platform & DevEx · Best AI Assistant for SPL", "NL→SPL có vòng tự sửa (self-critique)"),
        ("SOC Triage Copilot", "Security · Best MCP", "Điều tra alert tự động, đa bước có pivot"),
    ]
    _table(doc, rows)

    # 2. Kết quả kiểm thử
    h(doc, "2. Kết quả kiểm thử")
    doc.add_paragraph(
        "Lệnh: SKIP_STACK_CHECK=1 SKIP_SEED=1 python -m pytest -q  "
        "(testpaths = agents, tests, spl_copilot, soc_triage)."
    )
    p = doc.add_paragraph()
    run = p.add_run("KẾT QUẢ: 28 passed, 0 failed (keyless).")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x7F, 0x37)
    _table(doc, [
        ("Bộ test", "Số test", "Trạng thái"),
        ("Project chính (agents/, tests/)", "17", "PASS"),
        ("SPL Copilot (spl_copilot/tests)", "5", "PASS"),
        ("SOC Triage (soc_triage/tests)", "6", "PASS"),
        ("Tổng", "28", "PASS"),
    ])
    doc.add_paragraph(
        "Ghi chú: test tích hợp (tests/integration) và e2e (tests/e2e) cần stack thật "
        "(make up) + API key nên được loại khỏi lần chạy mặc định, đúng thiết kế của repo. "
        "Build web (npm run build) cũng PASS cho cả 2 route mới."
    )

    # 3. Bằng chứng giao diện
    h(doc, "3. Bằng chứng giao diện (screenshots)")
    captions = [
        ("01_dashboard.png", "Executive Pulse — dashboard điều hành (persona switcher, story cards, audio)."),
        ("02_spl_copilot.png", "SPL Copilot — agent tự phát hiện field `status` không có trong schema và remap sang `http_status`, chạy lại ra 3 dòng, kèm giải thích từng pipe."),
        ("03_soc_triage.png", "SOC Triage Copilot — điều tra 4 bước (pivot brute-force → đăng nhập thành công → export 1247 records), timeline tấn công, verdict CRITICAL + hành động containment."),
    ]
    for fname, cap in captions:
        img = SHOTS / fname
        if img.exists():
            doc.add_picture(str(img), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(cap)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c.runs:
            r.font.size = Pt(9)
            r.italic = True

    # 4. Tuân thủ tài liệu tiếng Anh
    h(doc, "4. Tuân thủ tài liệu (toàn bộ tiếng Anh)")
    doc.add_paragraph(
        "Đã quét toàn repo (*.md, *.py, *.ts, *.tsx, loại trừ .venv/node_modules/.next) "
        "tìm ký tự tiếng Việt có dấu: KHÔNG phát hiện. Toàn bộ README, comment, docstring "
        "của 2 project mới đều bằng tiếng Anh. (Riêng file audit này theo yêu cầu là tiếng Việt "
        "để review, không thuộc tài liệu kỹ thuật của repo.)"
    )

    # 5. Đề xuất nâng cấp theo yêu cầu cuộc thi
    h(doc, "5. Đề xuất nâng cấp theo yêu cầu cuộc thi")
    doc.add_paragraph("Ma trận năng lực Splunk bắt buộc vs. mức độ sử dụng (sau nâng cấp):")
    _table(doc, [
        ("Năng lực Splunk", "Executive Pulse", "SPL Copilot", "SOC Triage"),
        ("MCP Server", "Có", "Có (đường live)", "Có (inject searcher)"),
        ("AI Assistant for SPL", "Có (drill-down)", "Có (lõi)", "Chưa (đề xuất TB)"),
        ("Hosted Models (LLM)", "Có (narrative)", "Có (MỚI: explain + fix)", "Có (MỚI: narrative)"),
        ("AI Toolkit (MLTK)", "Có", "Không cần", "Có (MỚI: anomalydetection)"),
        ("AI agents", "Có (7 agent)", "Có", "Có"),
    ])
    p = doc.add_paragraph()
    r = p.add_run("Đã triển khai 3 nâng cấp ưu tiên CAO (giữ nguyên đường keyless qua cổng env + fallback offline):")
    r.bold = True
    _table(doc, [
        ("Trạng thái", "Sản phẩm", "Nâng cấp", "Giải được củng cố"),
        ("XONG", "SPL Copilot", "Hosted Models cho bước explain + LLM fallback sửa field khi alias/fuzzy thất bại", "Best Hosted Models"),
        ("XONG", "SOC Triage", "MLTK | anomalydetection để PHÁT HIỆN spike auth (bước 1, thay vì giả định alert)", "AI Toolkit + Observability"),
        ("XONG", "SOC Triage", "Hosted Models viết bản tường thuật incident (verdict vẫn deterministic)", "Best Hosted Models"),
    ])
    doc.add_paragraph("Đề xuất còn lại (chưa làm):")
    _table(doc, [
        ("Ưu tiên", "Sản phẩm", "Nâng cấp", "Giải được củng cố"),
        ("TB", "SPL Copilot", "MCPExecutor lấy schema thật qua `| fieldsummary` cho self-fix khi chạy live", "Best MCP"),
        ("TB", "SOC Triage", "Cho analyst hỏi follow-up bằng NL (tái dùng AI Assistant for SPL drill-down)", "Best AI Assistant for SPL"),
        ("TB", "Cả hai", "Demo live 1 query thật qua docker-compose có sẵn trong repo", "Tất cả"),
        ("Thấp", "SOC Triage", "Mở rộng playbook: data exfil, lateral movement; nối delivery/ cho approval", "Security"),
        ("Thấp", "Chung", "Trang landing nối 3 demo + quay video demo 2 phút mỗi sản phẩm", "Design / Impact"),
    ])

    # 6. Kết luận
    h(doc, "6. Kết luận")
    doc.add_paragraph(
        "Trạng thái hiện tại: 3 sản phẩm độc lập, 28/28 test pass keyless, web build sạch, "
        "tài liệu repo toàn tiếng Anh. 3 nâng cấp ưu tiên CAO đã được triển khai: Hosted Models "
        "trong SPL Copilot (explain + fix fallback), MLTK anomalydetection và Hosted Models "
        "narrative trong SOC Triage. Tất cả đều có cổng env + fallback offline nên giữ nguyên "
        "đường chạy keyless. Các đề xuất TB/Thấp còn lại (schema live qua fieldsummary, drill-down "
        "NL, demo live, video) là bước tiếp theo để tối đa hóa độ phủ giải."
    )

    doc.save(str(OUT))
    print(f"Saved {OUT}")


def _table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = val
            if i == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


if __name__ == "__main__":
    main()
